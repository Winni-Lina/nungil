import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 공통 헬퍼 ────────────────────────────────────────────────────────────────
def new_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(1.8)
        s.left_margin = s.right_margin = Cm(2.2)
    doc.styles['Normal'].font.name = '맑은 고딕'
    doc.styles['Normal'].font.size = Pt(10)
    return doc

C_TITLE  = RGBColor(0x1a,0x1a,0x2e); C_BLUE  = RGBColor(0x1a,0x5f,0x9e)
C_GREEN  = RGBColor(0x1a,0x7a,0x3c); C_RED   = RGBColor(0xcc,0x33,0x33)
C_GRAY   = RGBColor(0x66,0x66,0x66); C_ORANGE= RGBColor(0xaa,0x55,0x00)
C_PURPLE = RGBColor(0x6a,0x1a,0x9a)

def rn(p,text,bold=False,size=10,color=None):
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(size); r.font.name='맑은 고딕'
    if color: r.font.color.rgb=color
    return r

def hr(doc,col='BBBBBB'):
    p=doc.add_paragraph()
    pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
    b=OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6')
    b.set(qn('w:space'),'1');    b.set(qn('w:color'),col)
    pBdr.append(b); pPr.append(pBdr); p.paragraph_format.space_after=Pt(4)

def ctr(doc,text,bold=False,size=11,color=None):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rn(p,text,bold=bold,size=size,color=color); p.paragraph_format.space_after=Pt(2)

def sec(doc,num,title,color=None):
    p=doc.add_paragraph()
    rn(p,f'{num}.  {title}',bold=True,size=13,color=color or C_BLUE)
    p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)

def sub(doc,text,color=None,size=10.5):
    p=doc.add_paragraph()
    rn(p,text,bold=True,size=size,color=color or C_TITLE)
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)

def body(doc,text,color=None):
    p=doc.add_paragraph(); rn(p,text,color=color); p.paragraph_format.space_after=Pt(4)

def bul(doc,text,bold_prefix=None,color=None):
    p=doc.add_paragraph(style='List Bullet')
    if bold_prefix: rn(p,bold_prefix,bold=True,color=color)
    rn(p,text); p.paragraph_format.space_after=Pt(2)

def shade(cell,hex6):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'),hex6); tcPr.append(shd)

def tbl(doc,headers,rows,widths=None,hbg='1A5F9E'):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'
    for cell,h in zip(t.rows[0].cells,headers):
        shade(cell,hbg)
        r=cell.paragraphs[0].add_run(h)
        r.bold=True; r.font.name='맑은 고딕'; r.font.size=Pt(9.5)
        r.font.color.rgb=RGBColor(255,255,255)
    for ri,rd in enumerate(rows):
        row=t.add_row()
        for cell,val in zip(row.cells,rd):
            if ri%2==1: shade(cell,'EEF3FA')
            r2=cell.paragraphs[0].add_run(str(val))
            r2.font.name='맑은 고딕'; r2.font.size=Pt(9.5)
    if widths:
        for row2 in t.rows:
            for i,w in enumerate(widths): row2.cells[i].width=Cm(w)
    return t

def img(doc,path,width_cm=16.0):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path,width=Cm(width_cm))
    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(6)

def pg(doc):
    p=doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)


# ═══════════════════════════════════════════════════════════════════════════
# 문서 1: API 명세서
# ═══════════════════════════════════════════════════════════════════════════
doc1 = new_doc()
doc1.add_paragraph()
ctr(doc1,'눈길 (Nungil)  API 명세서',bold=True,size=20,color=C_TITLE)
ctr(doc1,'서버 엔드포인트 목록  |  2026. 06',size=11,color=C_GRAY)
doc1.add_paragraph(); hr(doc1)
body(doc1,'베이스 URL:  http://[서버IP]:8080    |    응답 형식: JSON    |    공통 응답: { "status": "SUCCESS"/"ERROR", "result": {...} }')
hr(doc1)

# ── 인증 ──
sec(doc1,'1','보호자 인증  —  /api/v1/guardian/auth')
tbl(doc1,
    ['메서드','경로','설명','주요 파라미터','응답'],
    [
        ('GET',   '/check-id',       'ID 중복 확인',        '?id=guardianId',                 'available: true/false'),
        ('POST',  '/signup',         '회원가입',            '{ id, pw, email, phone, name }', 'guardianId, name'),
        ('POST',  '/login',          '로그인',              '{ id, pw }',                     'guardianId, name'),
        ('POST',  '/reset-password', '비밀번호 재설정',     '{ id, email, newPw }',           'SUCCESS'),
        ('DELETE','/{id}',           '계정 삭제',           'Path: id',                       'SUCCESS'),
        ('PUT',   '/fcm-token',      'FCM 토큰 업데이트',   '{ id, fcmToken }',               'SUCCESS'),
    ],
    widths=[2,3.5,3,5,2.5]
)
doc1.add_paragraph()

# ── 사용자 관리 ──
sec(doc1,'2','사용자 관리  —  /api/v1/guardian')
tbl(doc1,
    ['메서드','경로','설명','주요 파라미터','응답'],
    [
        ('POST','  /users',                                '사용자 등록',          '{ guardianId, userName, userPhone }', 'idx(순번)'),
        ('GET', '  /settings/user/{guardianId}/{idx}',    '사용자 정보 조회',     'Path: guardianId, idx',              'NungilUserVO'),
        ('POST','  /settings/user/.../userinfo',          '사용자 정보 수정',     '{ userName, userPhone }',            'SUCCESS'),
        ('GET', '  /settings/user/.../whitelist',         '허용 과업 목록 조회',  'Path: guardianId, idx',              'whiteList 배열'),
        ('POST','  /settings/user/.../whitelist',         '허용 과업 추가',       '{ taskId }',                         'SUCCESS'),
        ('DELETE','  /settings/user/.../whitelist/{taskId}','허용 과업 제거',     'Path: guardianId, idx, taskId',      'SUCCESS'),
        ('GET', '  /tasks',                               '전체 과업 목록',       '없음',                               'List<TaskVO>'),
        ('GET', '  /tasks/search',                        '과업 검색',            '?item=검색어',                       'List<TaskVO>'),
    ],
    widths=[2,5,3,4,2]
)
doc1.add_paragraph()

pg(doc1)

# ── 일정 ──
sec(doc1,'3','일정 관리  —  /api/v1/guardian/schedule')
tbl(doc1,
    ['메서드','경로','설명','주요 파라미터','응답'],
    [
        ('POST',  '',                    '일정 생성',          '{ guardianId, idx, taskId, scheduledAt, location, specialNote }', 'scheduleId'),
        ('POST',  '/generate-steps',     'AI 단계 생성',       '{ taskId, taskName, location, specialNote }',  '["단계1","단계2",...]'),
        ('GET',   '',                    '일정 목록 조회',     '?guardianId=&idx=',                             'List<ScheduleVO>'),
        ('GET',   '/{guardianId}/{idx}', '사용자별 일정 조회', 'Path: guardianId, idx',                        'List<ScheduleVO>'),
        ('PUT',   '/{scheduleId}/complete','일정 완료 처리',   'Path: scheduleId',                              'SUCCESS'),
        ('PUT',   '/{scheduleId}/time',  '예약 시간 변경',     '{ scheduledAt }',                              'SUCCESS'),
        ('PATCH', '/{scheduleId}/status','상태 변경',          '{ status }',                                   'SUCCESS'),
        ('DELETE','/{scheduleId}',       '일정 삭제',          'Path: scheduleId',                             'SUCCESS'),
    ],
    widths=[2,4,3,5,2]
)
doc1.add_paragraph()

# ── 사용자 앱 ──
sec(doc1,'4','사용자 앱 API  —  /api/v1')
tbl(doc1,
    ['메서드','경로','설명','주요 파라미터','응답'],
    [
        ('POST',  '/guardian/login',          '보호자 로그인 (앱용)',  '{ guardianId, password }',       'guardianId, name'),
        ('GET',   '/user/{userId}/{userIdx}', '사용자 정보 조회',      'Path: userId, userIdx',          'NungilUserVO'),
        ('GET',   '/schedule',               '오늘 일정 조회',         '?userId=&userIdx=&date=',        'List<ScheduleVO>'),
        ('PATCH', '/schedule/{id}/complete', '일정 완료',              'Path: scheduleId',               'SUCCESS'),
        ('POST',  '/question/log',           '질문 횟수 기록',         '{ scheduleId }',                 'SUCCESS'),
        ('POST',  '/user/link',              '사용자-보호자 연결',     '{ guardianId, idx }',            'SUCCESS'),
        ('PUT',   '/user/fcm-token',         'FCM 토큰 갱신',          '{ userId, userIdx, fcmToken }',  'SUCCESS'),
        ('GET',   '/user/link/{guardianId}', '연결된 사용자 목록',     'Path: guardianId',               'List<NungilUserVO>'),
    ],
    widths=[2,4.5,3,5,1.5]
)
doc1.add_paragraph()

# ── AI 분석 ──
sec(doc1,'5','AI 대화  —  /api/v1/question')
tbl(doc1,
    ['메서드','경로','설명','파라미터','응답 필드'],
    [
        ('POST','/analyze','AI 대화 분석\n(핵심 API)',
         'mode: chat/schedule\ntextPrompt or voiceFile\nscheduleTitle, currentStep\nstepIndex, totalSteps\nspecialNote',
         'answer: 안내 문구\nstepComplete: 완료 여부\nsuggestedQuestions: 추천 질문\nphotoRequest: 카메라 요청'),
        ('POST','/summarize','일정 완료 요약',
         '{ scheduleTitle, historyJson }',
         'message: 완료 요약 문구'),
    ],
    widths=[2,3,3,5,4]
)

doc1.add_paragraph(); hr(doc1)
ctr(doc1,'눈길 팀  |  한이음 ICT 멘토링 클럽  |  2026. 06',size=9,color=C_GRAY)
doc1.save(r'C:\Users\User\Desktop\눈길_API명세서.docx')
print('API 명세서 저장 완료')


# ═══════════════════════════════════════════════════════════════════════════
# 문서 2: 사용자 매뉴얼
# ═══════════════════════════════════════════════════════════════════════════
doc2 = new_doc()
doc2.add_paragraph()
ctr(doc2,'눈길 (Nungil)',bold=True,size=22,color=C_TITLE)
ctr(doc2,'사용자 앱 매뉴얼',bold=True,size=16,color=C_TITLE)
ctr(doc2,'"똘똘이"와 함께하는 일상 도우미',size=12,color=C_GRAY)
doc2.add_paragraph(); hr(doc2)

# 시작 방법
sec(doc2,'1','시작하는 방법')
body(doc2,'앱을 처음 열면 보호자 분이 미리 연결해 뒀을 거예요. 연결되어 있으면 바로 사용할 수 있어요.')
tbl(doc2,
    ['순서','방법'],
    [
        ('1','스마트폰에서 눈길 앱 아이콘을 눌러요'),
        ('2','"똘똘이" 친구 화면이 뜨면 준비된 거예요'),
        ('3','"똘똘아" 라고 부르거나, 화면을 눌러서 시작해요'),
    ],
    widths=[2,14]
)
doc2.add_paragraph()

# 똘똘이 부르는 법
sec(doc2,'2','똘똘이 부르는 방법')
body(doc2,'똘똘이는 항상 듣고 있어요. 이렇게 불러보세요!')
tbl(doc2,
    ['방법','설명'],
    [
        ('"똘똘아" 라고 말하기', '폰을 들고 "똘똘아" 라고 말하면 바로 반응해요'),
        ('화면 마이크 버튼 누르기', '화면 아래 동그란 마이크 버튼을 눌러도 돼요'),
        ('똘똘이가 말하는 중에 누르기', '말하는 도중 버튼을 누르면 멈추고 들어줘요'),
    ],
    widths=[5,11]
)
doc2.add_paragraph()

# 일정 수행
sec(doc2,'3','일정 따라 하기')
body(doc2,'알람이 울리면 똘똘이가 오늘 해야 할 일을 하나씩 알려줘요.')
tbl(doc2,
    ['상황','내가 할 것'],
    [
        ('알람이 울려요',               '알람을 누르면 똘똘이가 첫 번째 할 일을 알려줘요'),
        ('똘똘이가 할 일을 알려줘요',   '잘 듣고, 그 일을 해보세요'),
        ('다 했어요!',                  '"했어요", "응", "다 했어" 라고 말해요\n→ 똘똘이가 다음 할 일을 알려줘요'),
        ('모르겠어요',                  '"어떻게 해요?" 라고 물어보세요\n→ 똘똘이가 더 쉽게 설명해줘요'),
        ('힘들어요',                    '"힘들어요", "못 하겠어요" 라고 말해도 돼요\n→ 똘똘이가 천천히 도와줄게요'),
        ('모두 다 했어요!',             '똘똘이가 "정말 잘했어요!" 라고 말해줘요\n→ 보호자분도 알게 돼요'),
    ],
    widths=[4.5,11.5]
)
doc2.add_paragraph()

# 버튼 설명
sec(doc2,'4','화면 버튼 설명')
tbl(doc2,
    ['버튼','하는 일'],
    [
        ('마이크 버튼 (동그란 버튼)', '누르면 말하기 시작, 다시 누르면 멈춰요'),
        ('"했어요" 버튼',             '말하기 어려울 때 눌러도 완료돼요'),
        ('"모르겠어요" 버튼',         '뭘 해야 할지 모를 때 눌러요 — 똘똘이가 다시 알려줘요'),
        ('"다음 단계로" 버튼',        '현재 할 일이 끝났을 때 나타나요'),
    ],
    widths=[5,11]
)
doc2.add_paragraph()

# 자유 대화
sec(doc2,'5','궁금한 거 물어보기')
body(doc2,'일정 말고도 궁금한 게 있으면 언제든지 물어볼 수 있어요!')
tbl(doc2,
    ['물어볼 수 있는 것','예시'],
    [
        ('주변 물건이 뭔지',   '"이게 뭐야?" — 카메라로 사진 찍으면 알려줘요'),
        ('오늘 날씨',          '"오늘 날씨 어때?" 라고 말해요'),
        ('심심할 때',          '"심심해" 라고 말해요 — 똘똘이가 같이 놀아줘요'),
        ('뭐든 궁금한 것',     '그냥 말하면 돼요! 똘똘이가 항상 들어줘요'),
    ],
    widths=[4.5,11.5]
)
doc2.add_paragraph()

# 주의사항
sec(doc2,'6','이럴 때는 어른에게 말해요')
tbl(doc2,
    ['상황','할 일'],
    [
        ('몸이 아파요',             '가까운 어른을 불러요. 똘똘이한테 말하면 도움말을 알려줘요'),
        ('앱이 안 켜져요',          '보호자분께 말씀드려요'),
        ('똘똘이가 이상한 말을 해요','보호자분께 말씀드려요'),
    ],
    widths=[4.5,11.5]
)

doc2.add_paragraph(); hr(doc2)
ctr(doc2,'눈길 팀  |  한이음 ICT 멘토링 클럽  |  2026. 06',size=9,color=C_GRAY)
doc2.save(r'C:\Users\User\Desktop\눈길_사용자매뉴얼.docx')
print('사용자 매뉴얼 저장 완료')


# ═══════════════════════════════════════════════════════════════════════════
# 문서 3: ERD 문서 (그림 + 테이블 설명)
# ═══════════════════════════════════════════════════════════════════════════
doc3 = new_doc()
doc3.add_paragraph()
ctr(doc3,'눈길 (Nungil)  DB 구조 (ERD)',bold=True,size=20,color=C_TITLE)
ctr(doc3,'Entity Relationship Diagram  |  2026. 06',size=11,color=C_GRAY)
doc3.add_paragraph(); hr(doc3)

img(doc3, r'C:\Users\User\Desktop\diagram_erd.png', width_cm=16.0)

sec(doc3,'1','테이블 구조')

sub(doc3,'GUARDIAN  —  보호자 계정')
tbl(doc3,
    ['컬럼','타입','설명'],
    [
        ('id (PK)',    'VARCHAR2', '보호자 로그인 ID'),
        ('pw',        'VARCHAR2', '비밀번호'),
        ('email',     'VARCHAR2', '이메일'),
        ('phone',     'VARCHAR2', '전화번호'),
        ('name',      'VARCHAR2', '이름'),
        ('fcmToken',  'VARCHAR2', '푸시 알림 토큰'),
        ('regdate',   'TIMESTAMP','가입일시'),
    ],
    widths=[4,3,9]
)
doc3.add_paragraph()

sub(doc3,'NUNGIL_USER  —  사용자 (장애인)')
tbl(doc3,
    ['컬럼','타입','설명'],
    [
        ('id (PK/FK)', 'VARCHAR2', '보호자 ID (GUARDIAN.id 참조)'),
        ('idx (PK)',   'NUMBER',   '보호자당 사용자 순번 (1부터)'),
        ('userName',   'VARCHAR2', '사용자 이름'),
        ('userPhone',  'VARCHAR2', '사용자 전화번호'),
        ('specialNote','VARCHAR2', '특이사항 (예: 큰 소리 무서워함)'),
        ('whiteList',  'VARCHAR2', '허용된 과업 ID 목록 (쉼표 구분, 예: "1,2,3")'),
        ('fcmToken',   'VARCHAR2', '사용자 기기 푸시 알림 토큰'),
    ],
    widths=[4,3,9]
)
doc3.add_paragraph()

sub(doc3,'TASK  —  과업 (일상 활동 항목)')
tbl(doc3,
    ['컬럼','타입','설명'],
    [
        ('taskId (PK)',    'NUMBER',  '과업 고유 번호'),
        ('name',          'VARCHAR2','과업 이름 (예: 빨래하기)'),
        ('process',       'CLOB',    '과업 기본 단계 가이드 텍스트'),
        ('guardianId (FK)','VARCHAR2','등록한 보호자 ID'),
    ],
    widths=[4,3,9]
)
doc3.add_paragraph()

sub(doc3,'SCHEDULE  —  일정 (과업 실행 예약)')
tbl(doc3,
    ['컬럼','타입','설명'],
    [
        ('scheduleId (PK)', 'NUMBER',   '일정 고유 번호'),
        ('taskId (FK)',      'NUMBER',   '연결된 과업 (TASK.taskId 참조)'),
        ('id (FK)',          'VARCHAR2', '담당 보호자 ID'),
        ('idx (FK)',         'NUMBER',   '담당 사용자 순번'),
        ('status',           'VARCHAR2', '상태 (pending / in_progress / completed)'),
        ('scheduledAt',      'TIMESTAMP','예약된 일정 시각'),
        ('createdAt',        'TIMESTAMP','생성 시각'),
        ('successAt',        'TIMESTAMP','완료 시각'),
        ('location',         'VARCHAR2', '장소 (세탁실·부엌·거실·화장실·방·기타)'),
        ('specialNote',      'VARCHAR2', '이번 일정 특이사항'),
        ('customSteps',      'CLOB',     'AI가 생성·보호자가 확정한 맞춤 단계 목록 (JSON)'),
        ('questionCount',    'NUMBER',   '사용자가 질문한 횟수'),
        ('lastQuestionAt',   'TIMESTAMP','마지막 질문 시각'),
    ],
    widths=[4,3,9]
)
doc3.add_paragraph()

sec(doc3,'2','테이블 관계')
tbl(doc3,
    ['관계','설명'],
    [
        ('GUARDIAN  1 : N  NUNGIL_USER', '한 보호자가 여러 사용자를 관리할 수 있음 (최대 3명)'),
        ('GUARDIAN  1 : N  TASK',        '한 보호자가 여러 과업을 등록할 수 있음'),
        ('TASK  1 : N  SCHEDULE',        '하나의 과업으로 여러 일정을 만들 수 있음 (반복 일정)'),
        ('GUARDIAN + NUNGIL_USER  →  SCHEDULE', '일정은 특정 보호자의 특정 사용자에게 배정됨'),
    ],
    widths=[7,9]
)

doc3.add_paragraph(); hr(doc3)
ctr(doc3,'눈길 팀  |  한이음 ICT 멘토링 클럽  |  2026. 06',size=9,color=C_GRAY)
doc3.save(r'C:\Users\User\Desktop\눈길_DB구조(ERD).docx')
print('ERD 문서 저장 완료')
