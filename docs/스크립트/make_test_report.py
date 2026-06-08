import sys, datetime
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(1.8)
    s.left_margin = s.right_margin = Cm(2.2)
doc.styles['Normal'].font.name = '맑은 고딕'
doc.styles['Normal'].font.size = Pt(10)

C_TITLE  = RGBColor(0x1a, 0x1a, 0x2e)
C_BLUE   = RGBColor(0x1a, 0x5f, 0x9e)
C_GREEN  = RGBColor(0x1a, 0x7a, 0x3c)
C_RED    = RGBColor(0xcc, 0x33, 0x33)
C_GRAY   = RGBColor(0x66, 0x66, 0x66)
C_ORANGE = RGBColor(0xaa, 0x55, 0x00)

def rn(p, text, bold=False, size=10, color=None):
    r = p.add_run(text)
    r.bold=bold; r.font.size=Pt(size); r.font.name='맑은 고딕'
    if color: r.font.color.rgb=color
    return r

def hr(doc, col='BBBBBB'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6')
    b.set(qn('w:space'),'1');    b.set(qn('w:color'),col)
    pBdr.append(b); pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)

def center(doc, text, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(p, text, bold=bold, size=size, color=color)
    p.paragraph_format.space_after = Pt(2)

def sec(doc, num, title):
    p = doc.add_paragraph()
    rn(p, f'{num}.  {title}', bold=True, size=13, color=C_BLUE)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)

def sub(doc, text, color=None):
    p = doc.add_paragraph()
    rn(p, text, bold=True, size=10.5, color=color or C_TITLE)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)

def body(doc, text, color=None):
    p = doc.add_paragraph()
    rn(p, text, color=color)
    p.paragraph_format.space_after = Pt(4)

def shade(cell, hex6):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'),hex6); tcPr.append(shd)

def tbl(doc, headers, rows, widths=None, hbg='1A5F9E'):
    t = doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'
    for cell,h in zip(t.rows[0].cells, headers):
        shade(cell, hbg)
        r=cell.paragraphs[0].add_run(h)
        r.bold=True; r.font.name='맑은 고딕'; r.font.size=Pt(9.5)
        r.font.color.rgb=RGBColor(255,255,255)
    for ri,rd in enumerate(rows):
        row=t.add_row()
        for cell,val in zip(row.cells, rd):
            if ri%2==1: shade(cell,'EEF3FA')
            r2=cell.paragraphs[0].add_run(str(val))
            r2.font.name='맑은 고딕'; r2.font.size=Pt(9.5)
    if widths:
        for row2 in t.rows:
            for i,w in enumerate(widths): row2.cells[i].width=Cm(w)
    return t

def img(doc, path, width_cm=15):
    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width_cm))
    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(6)

def pgbreak(doc):
    p=doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)

# ── 상태 칸 색 헬퍼 ─────────────────────────────────────────────────────────
def pass_row(row_cells, passed=True):
    bg = '1A7A3C' if passed else 'CC3333'
    shade(row_cells[-1], bg)
    for run in row_cells[-1].paragraphs[0].runs:
        run.font.color.rgb = RGBColor(255,255,255)
        run.bold = True

# ════════════════════ 표지 ════════════════════
doc.add_paragraph()
center(doc, '눈길 (Nungil)', bold=True, size=22, color=C_TITLE)
center(doc, '테스트 결과 보고서', bold=True, size=16, color=C_TITLE)
center(doc, '2026년 6월 8일  |  한이음 ICT 멘토링 클럽', size=11, color=C_GRAY)
doc.add_paragraph()
hr(doc)

# 요약 박스
p = doc.add_paragraph()
rn(p, '실행 일시:  ', bold=True, size=11)
rn(p, '2026-06-08  (mvn test)', size=11)
p.paragraph_format.space_after = Pt(3)
p = doc.add_paragraph()
rn(p, '최종 결과:  ', bold=True, size=11)
rn(p, '33개 실행  /  30개 통과  /  0개 실패  /  3개 skip (실제 API 필요)', size=11, color=C_GREEN)
p.paragraph_format.space_after = Pt(8)
hr(doc)

# ════════════════════ 1. 전체 요약 ════════════════════
sec(doc, '1', '전체 요약')

tbl(doc,
    ['테스트 파일', '대상 기능', '총', '통과', '실패', 'skip'],
    [
        ('NungilAnalyzeE2ETest',      '앱↔서버 전체 흐름 (E2E)',    '8',  '8',  '0', '0'),
        ('AnalysisOrchestratorTest',  'AI 파이프라인 핵심 로직',    '4',  '4',  '0', '0'),
        ('StepGenerationServiceTest', 'AI 단계 생성·파싱 로직',     '4',  '4',  '0', '0'),
        ('GuardianServiceTest',       '보호자 회원·로그인 서비스',  '10', '10', '0', '0'),
        ('ScheduleServiceTest',       '일정 도메인 서비스',         '4',  '4',  '0', '0'),
        ('GeminiLiveIntegrationTest', '실제 Gemini API 연동',       '3',  '0',  '0', '3'),
        ('합계', '', '33', '30', '0', '3'),
    ],
    widths=[5, 5.5, 1, 1.5, 1.5, 1.5]
)
doc.add_paragraph()

# ════════════════════ 2. E2E 테스트 ════════════════════
pgbreak(doc)
sec(doc, '2', 'E2E 테스트  —  NungilAnalyzeE2ETest')

body(doc,
    '앱이 서버로 요청을 보냈을 때, 서버가 올바른 응답을 돌려주는지 전체 흐름을 확인합니다. '
    'Gemini AI는 가짜 응답으로 대체해서 API 키 없이도 실행됩니다.')

tbl(doc,
    ['번호', '시나리오', '요청 내용', '기대 결과', '결과'],
    [
        ('E2E-01', '채팅 — 일반 질문',
         'mode=chat\ntextPrompt="오늘 뭐 먹어?"',
         'status=SUCCESS\nstepComplete=false\nanswer 있음',
         '통과'),
        ('E2E-02', '채팅 — 사물 질문\n→ 카메라 요청',
         'mode=chat\ntextPrompt="이게 뭐야?"',
         'photoRequest=true',
         '통과'),
        ('E2E-03', '채팅 — 응답 필드\n완전성 검증',
         'mode=chat\ntextPrompt="안녕"',
         'answer·stepComplete·\nsuggestedQuestions·\nphotoRequest 모두 존재',
         '통과'),
        ('E2E-04', '일정 — 첫 단계 시작',
         'mode=schedule\nstepIndex=0\ntextPrompt="일정 시작"',
         'stepComplete=false\nanswer 있음',
         '통과'),
        ('E2E-05', '일정 — 완료 신호\n"다 했어"',
         'mode=schedule\ntextPrompt="다 했어"',
         'stepComplete=true\n칭찬 answer',
         '통과'),
        ('E2E-06', '일정 — 방법 질문\n"어떻게 해?"',
         'mode=schedule\ntextPrompt="어떻게 해?"',
         'stepComplete=false\n설명 answer',
         '통과'),
        ('E2E-07', '일정 — 힘들어함\n"못 하겠어"',
         'mode=schedule\ntextPrompt="못 하겠어"',
         'stepComplete=false\n응원 answer',
         '통과'),
        ('E2E-08', '일정 — 음성 파일\n첨부 경로',
         'mode=schedule\nvoiceFile=voice.wav',
         'status=SUCCESS',
         '통과'),
    ],
    widths=[1.5, 3, 4, 4.5, 2]
)

# 결과 열 초록색
t = doc.tables[-1]
for row in list(t.rows)[1:]:
    cells = row.cells
    val = cells[-1].paragraphs[0].runs[0].text if cells[-1].paragraphs[0].runs else ''
    if val == '통과': pass_row(cells)

doc.add_paragraph()

# ════════════════════ 3. 단위 테스트 ════════════════════
pgbreak(doc)
sec(doc, '3', '단위 테스트  —  기능별 상세')

# 3-1 AI 파이프라인
sub(doc, '3-1  AI 파이프라인  —  AnalysisOrchestratorTest')
body(doc, 'AI에게 보내는 프롬프트가 올바르게 만들어지는지, 응답 파싱이 정상인지 확인합니다.')
tbl(doc,
    ['번호', '테스트 내용', '확인 항목', '결과'],
    [
        ('U-01', 'buildChatPrompt — system/user 2개 반환',
         'system에 "똘똘이" 포함\nuser에 사용자 발화 포함', '통과'),
        ('U-02', 'buildSchedulePrompt — 단계 안내 구조',
         '"앱이 결정" 문구 포함\nuser에 일정명·단계명 포함', '통과'),
        ('U-03', 'safeHistory — 빈 대화 기록 처리',
         'null·""·"[]" → "(대화 없음)" 반환', '통과'),
        ('U-04', 'execute — Gemini 1번만 호출',
         'Gemini 호출 횟수=1\nanswer·userId 정상 반환', '통과'),
    ],
    widths=[1.5, 5, 6, 2]
)
_t = doc.tables[-1]
for row in list(_t.rows)[1:]:
    cells=row.cells
    val=cells[-1].paragraphs[0].runs[0].text if cells[-1].paragraphs[0].runs else ''
    if val=='통과': pass_row(cells)

doc.add_paragraph()

# 3-2 단계 생성
sub(doc, '3-2  AI 단계 생성  —  StepGenerationServiceTest')
body(doc, 'Gemini가 단계를 생성했을 때 다양한 응답 형식을 올바르게 처리하는지 확인합니다.')
tbl(doc,
    ['번호', '테스트 내용', '확인 항목', '결과'],
    [
        ('U-05', 'JSON 펜스로 감싼 응답 파싱',
         '```json [...] ``` → 배열 정상 파싱', '통과'),
        ('U-06', '줄바꿈 텍스트 fallback 파싱',
         'JSON 아닌 텍스트 → taskProcess 기반 파싱', '통과'),
        ('U-07', '빈 응답 처리',
         'Gemini가 빈 문자열 반환 → 빈 리스트', '통과'),
        ('U-08', 'JSON 배열 문자열 정상 파싱',
         '["단계1","단계2"] → List<String>', '통과'),
    ],
    widths=[1.5, 5, 6, 2]
)
_t=doc.tables[-1]
for row in list(_t.rows)[1:]:
    cells=row.cells
    val=cells[-1].paragraphs[0].runs[0].text if cells[-1].paragraphs[0].runs else ''
    if val=='통과': pass_row(cells)

doc.add_paragraph()

# 3-3 보호자 서비스
sub(doc, '3-3  보호자 서비스  —  GuardianServiceTest')
body(doc, '보호자 회원가입, 로그인, 비밀번호 변경의 정상·오류 케이스를 확인합니다.')
tbl(doc,
    ['번호', '테스트 내용', '확인 항목', '결과'],
    [
        ('U-09',  '회원가입 성공', 'DB insert 1회 호출', '통과'),
        ('U-10',  '회원가입 실패 — ID 중복', 'IllegalArgumentException (ID_EXISTS)', '통과'),
        ('U-11',  '회원가입 실패 — 이메일 중복', 'IllegalArgumentException (EMAIL_EXISTS)', '통과'),
        ('U-12',  'ID 중복 확인 — 사용 가능', 'true 반환', '통과'),
        ('U-13',  'ID 중복 확인 — 이미 존재', 'false 반환', '통과'),
        ('U-14',  '로그인 실패 — 없는 ID', 'INVALID_CREDENTIALS 예외', '통과'),
        ('U-15',  '로그인 실패 — 비밀번호 불일치', 'INVALID_CREDENTIALS 예외', '통과'),
        ('U-16',  '로그인 성공', 'GuardianVO 반환', '통과'),
        ('U-17',  '비밀번호 재설정 실패 — 없는 사용자', 'USER_NOT_FOUND 예외', '통과'),
        ('U-18',  '비밀번호 재설정 실패 — 이메일 불일치', 'EMAIL_MISMATCH 예외', '통과'),
    ],
    widths=[1.5, 5, 6, 2]
)
_t=doc.tables[-1]
for row in list(_t.rows)[1:]:
    cells=row.cells
    val=cells[-1].paragraphs[0].runs[0].text if cells[-1].paragraphs[0].runs else ''
    if val=='통과': pass_row(cells)

doc.add_paragraph()

# 3-4 일정 서비스
sub(doc, '3-4  일정 서비스  —  ScheduleServiceTest')
body(doc, '일정 생성·완료·수정·삭제가 DB까지 올바르게 전달되는지 확인합니다.')
tbl(doc,
    ['번호', '테스트 내용', '확인 항목', '결과'],
    [
        ('U-19', '일정 생성', 'Mapper.insert 호출됨', '통과'),
        ('U-20', '일정 완료', 'Mapper.updateSuccessAt 호출됨', '통과'),
        ('U-21', '예약 시간 변경', '변경 시간 그대로 전달됨', '통과'),
        ('U-22', '일정 삭제', 'Mapper.deleteById 호출됨', '통과'),
    ],
    widths=[1.5, 5, 6, 2]
)
_t=doc.tables[-1]
for row in list(_t.rows)[1:]:
    cells=row.cells
    val=cells[-1].paragraphs[0].runs[0].text if cells[-1].paragraphs[0].runs else ''
    if val=='통과': pass_row(cells)

doc.add_paragraph()

# 3-5 실제 연동 (skip)
sub(doc, '3-5  실제 AI 연동  —  GeminiLiveIntegrationTest  (실행 조건부)')
body(doc,
    '실제 Gemini API 키가 있을 때만 실행하는 테스트입니다. '
    '일반 빌드에서는 자동으로 건너뛰고, 별도 명령으로 실행합니다.')
tbl(doc,
    ['번호', '테스트 내용', '실행 조건', '결과'],
    [
        ('L-01', '채팅 프롬프트 실제 호출 → JSON 파싱 가능', 'API 키 필요', 'skip'),
        ('L-02', '일정 단계 생성 실제 호출 → 빈 배열 아님', 'API 키 필요', 'skip'),
        ('L-03', '일정 완료 요약 실제 호출 → 문자열 반환',  'API 키 필요', 'skip'),
    ],
    widths=[1.5, 7, 3, 2]
)
_t=doc.tables[-1]
for row in list(_t.rows)[1:]:
    cells=row.cells
    val=cells[-1].paragraphs[0].runs[0].text if cells[-1].paragraphs[0].runs else ''
    shade(cells[-1], 'AAAAAA')
    for r in cells[-1].paragraphs[0].runs:
        r.font.color.rgb=RGBColor(255,255,255); r.bold=True

doc.add_paragraph()

# ════════════════════ 4. 커버리지 ════════════════════
pgbreak(doc)
sec(doc, '4', '코드 커버리지  (JaCoCo)')

body(doc,
    '전체 서버 코드 중 테스트로 실행된 비율입니다. '
    '웹 요청 처리 코드(컨트롤러)는 단위 테스트로 측정하기 어려운 구조여서 전체 수치가 낮습니다. '
    '비즈니스 핵심 로직(AI 파이프라인)은 52%를 확보했습니다.')

tbl(doc,
    ['측정 대상', '라인 커버리지', '내용'],
    [
        ('전체 서버 (평균)',              '12%  (169 / 1,403)',  '컨트롤러·인프라 계층 포함 전체'),
        ('AI 파이프라인 (infrastructure.google)', '52%  (100 / 190)', '핵심 비즈니스 로직'),
        ('보호자 도메인 (domain.guardian)',        '51%  (31 / 60)',  '회원·인증 로직'),
        ('일정 도메인 (domain.schedule)',           '32%  (21 / 64)',  '일정 상태 관리'),
        ('API 서버 (api.nungil)',                   '5%   (17 / 329)', 'E2E 테스트 추가로 신규 측정'),
        ('API 컨트롤러 계층 (기타)',                '0%',              '추후 통합 테스트 추가 예정'),
    ],
    widths=[5.5, 4, 6.5]
)
doc.add_paragraph()

img(doc, r'C:\Users\User\Desktop\diagram_coverage.png', width_cm=13.5)

doc.add_paragraph()
hr(doc)
center(doc, '눈길 팀  |  한이음 ICT 멘토링 클럽  |  2026. 06. 08', size=9, color=C_GRAY)

out = r'C:\Users\User\Desktop\눈길_테스트결과보고서.docx'
doc.save(out)
print(f'저장 완료: {out}')
