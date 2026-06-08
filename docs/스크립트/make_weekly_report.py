import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(1.8)
    s.left_margin = s.right_margin = Cm(2.2)

doc.styles['Normal'].font.name = '맑은 고딕'
doc.styles['Normal'].font.size = Pt(10)

C_TITLE  = RGBColor(0x1a, 0x1a, 0x2e)
C_BLUE   = RGBColor(0x1a, 0x5f, 0x9e)
C_GREEN  = RGBColor(0x1a, 0x7a, 0x3c)
C_RED    = RGBColor(0xcc, 0x33, 0x33)
C_ORANGE = RGBColor(0xaa, 0x55, 0x00)
C_GRAY   = RGBColor(0x66, 0x66, 0x66)
C_PURPLE = RGBColor(0x6a, 0x1a, 0x9a)

def rn(p, text, bold=False, size=10, color=None):
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size); r.font.name = '맑은 고딕'
    if color: r.font.color.rgb = color
    return r

def hr(doc, col='BBBBBB'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6')
    b.set(qn('w:space'),'1');    b.set(qn('w:color'),col)
    pBdr.append(b); pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)

def center(doc, text, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(p, text, bold=bold, size=size, color=color)
    p.paragraph_format.space_after = Pt(2)

def sec(doc, num, title):
    p = doc.add_paragraph()
    rn(p, f'{num}.  {title}', bold=True, size=13, color=C_BLUE)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)

def sub(doc, text, color=None):
    p = doc.add_paragraph()
    rn(p, text, bold=True, size=10.5, color=color or C_TITLE)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(3)

def body(doc, text, color=None):
    p = doc.add_paragraph()
    rn(p, text, color=color)
    p.paragraph_format.space_after = Pt(4)

def shade(cell, hex6):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'),hex6); tcPr.append(shd)

def tbl(doc, headers, rows, widths=None, hbg='1A5F9E'):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
    for cell, h in zip(t.rows[0].cells, headers):
        shade(cell, hbg)
        r = cell.paragraphs[0].add_run(h)
        r.bold=True; r.font.name='맑은 고딕'; r.font.size=Pt(9.5)
        r.font.color.rgb=RGBColor(255,255,255)
    for ri, rd in enumerate(rows):
        row = t.add_row()
        for cell, val in zip(row.cells, rd):
            if ri%2==1: shade(cell,'EEF3FA')
            r2 = cell.paragraphs[0].add_run(str(val))
            r2.font.name='맑은 고딕'; r2.font.size=Pt(9.5)
    if widths:
        for row2 in t.rows:
            for i,w in enumerate(widths): row2.cells[i].width=Cm(w)
    return t

def img(doc, path, width_cm=16.5):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width_cm))
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)

def pgbreak(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

# ════════════════════ 표지 ════════════════════
doc.add_paragraph()
center(doc, '눈길 (Nungil)', bold=True, size=22, color=C_TITLE)
center(doc, '주간 개발 현황 보고서', bold=True, size=15, color=C_TITLE)
center(doc, '2026년 6월 1주차  |  한이음 ICT 멘토링 클럽', size=11, color=C_GRAY)
doc.add_paragraph()
hr(doc)

# 한눈에 보는 현재 상태
p = doc.add_paragraph()
rn(p, '현재 상태:  ', bold=True, size=12)
rn(p, '핵심 기능 개발 완료. 남은 기간 UI 다듬기 + 발표 준비.', size=12, color=C_GREEN)
p.paragraph_format.space_after = Pt(3)

p = doc.add_paragraph()
rn(p, '이번 주 완료:  ', bold=True, size=11)
rn(p, 'AI 구조 전면 개편  /  전체 서비스 통합  /  테스트 30개 통과  /  결제 전환', size=11)
p.paragraph_format.space_after = Pt(3)

p = doc.add_paragraph()
rn(p, '마감까지:  ', bold=True, size=11)
rn(p, '6월 15일 (일)  —  7일 남음', size=11, color=C_RED)
p.paragraph_format.space_after = Pt(8)
hr(doc)

# ════════════════════ 1. 개발 완료 현황 ════════════════════
sec(doc, '1', '개발 완료 현황')

body(doc,
    '눈길의 핵심 기능은 전부 구현되어 동작하고 있습니다. '
    '보호자 앱과 사용자 앱이 서버를 통해 연결되고, AI가 음성으로 일정을 안내하며, '
    '사용자의 말에 반응하는 전체 흐름이 실제 기기에서 작동합니다.')

tbl(doc,
    ['기능', '상태', '비고'],
    [
        ('보호자 앱 — 회원가입·로그인',                 '완료', ''),
        ('보호자 앱 — 일정 등록 및 AI 단계 자동 생성', '완료', '등록 시 1회 생성·저장'),
        ('보호자 앱 — 생성 단계 검토·편집',            '완료', '보호자가 직접 수정 가능'),
        ('사용자 앱 — 웨이크워드 "똘똘아" 감지',       '완료', '인터넷 없이 동작'),
        ('사용자 앱 — 단계별 일정 수행 (음성)',         '완료', 'STT + TTS 연동'),
        ('사용자 앱 — 완료 키워드 감지 후 다음 단계',  '완료', '"했어" / "응" / "끝" 등'),
        ('사용자 앱 — 사진 기반 질문',                  '완료', '카메라 → AI 분석'),
        ('서버 — Google Gemini 2.5 AI',                '완료', '결제 계정 전환 완료'),
        ('서버 — 음성 인식 (Google STT)',               '완료', '서비스 계정 키 인증'),
        ('서버 — 푸시 알림 (FCM)',                      '완료', '일정 시작 알림'),
        ('서버 — 일정 완료 후 AI 요약',                 '완료', '보호자가 결과 확인 가능'),
        ('단위 테스트 22개 + E2E 테스트 8개',           '완료', '30개 전부 통과, 0개 실패'),
    ],
    widths=[6.5, 2, 7.5]
)
doc.add_paragraph()

pgbreak(doc)

# ════════════════════ 2. AI 구조 개편 ════════════════════
sec(doc, '2', 'AI 구조 개편 (이번 주 핵심 작업)')

body(doc,
    '가장 큰 문제였던 "AI가 실행마다 단계를 새로 만드는 구조"를 이번 주에 완전히 바꿨습니다. '
    '등록 시 단계를 한 번만 만들고, 실행 시에는 저장된 단계를 읽어서 안내만 합니다. '
    '완료 판단도 AI 대신 앱이 직접 합니다.')

img(doc, r'C:\Users\User\Desktop\diagram_before_after.png', width_cm=16.5)

pgbreak(doc)

sub(doc, '똘똘이가 사용자 말에 반응하는 방식')
body(doc,
    '사용자가 어떤 말을 해도 4가지로 구분해서 반응합니다. '
    '일정 완료 신호면 칭찬하고 넘어가고, 질문이면 쉽게 설명하고, '
    '힘들다고 하면 재촉하지 않고 응원합니다.')

img(doc, r'C:\Users\User\Desktop\diagram_4class.png', width_cm=16.5)

pgbreak(doc)

# ════════════════════ 3. 테스트 결과 ════════════════════
sec(doc, '3', '테스트 결과')

body(doc,
    '기능이 망가지지 않았는지 자동으로 확인하는 테스트를 두 종류로 작성했습니다. '
    '30개 전부 통과했습니다.')

tbl(doc,
    ['종류', '뭘 확인하는가', '개수', '결과'],
    [
        ('단위 테스트',
         'AI 프롬프트 생성, 완료 판단 로직, 도메인 서비스\n등 기능 하나하나 개별 확인',
         '22개', '전부 통과'),
        ('E2E 테스트',
         '앱이 서버에 요청을 보내고 응답이 올바르게 오는지\n실제 사용 시나리오 전체 흐름 확인\n예) "다 했어" → 완료 응답 / "못 하겠어" → 응원 응답',
         '8개', '전부 통과'),
        ('AI 실제 연동',
         '실제 Gemini API 호출 테스트\n(API 키 필요 — 별도 실행)',
         '3개', '별도 실행'),
    ],
    widths=[2.5, 9, 2, 2.5]
)
doc.add_paragraph()

img(doc, r'C:\Users\User\Desktop\diagram_coverage.png', width_cm=13.5)

pgbreak(doc)

# ════════════════════ 4. 현재 AI 불안정 원인과 개선 계획 ════════════════════
sec(doc, '4', '현재 AI가 완전하지 않은 이유와 개선 계획')

body(doc,
    '전체 흐름은 동작하지만, AI 응답이 아직 매끄럽지 않은 경우가 있습니다. '
    '아래 원인을 파악했고, 마무리 기간에 순서대로 개선합니다.')

tbl(doc,
    ['문제', '원인', '개선 방법', '시기'],
    [
        ('안내 문구가 너무 김',
         'AI에게 "짧게"라고만 했고\n글자 수 제한을 명확히 안 걸었음',
         'AI 지시에 "2문장, 30자 이내"\n조건 추가',
         '6/9~10'),
        ('단계 넘어갈 때 가끔 혼동',
         '이전 대화가 쌓이면\nAI가 몇 번째 단계인지 헷갈려 함',
         '단계 바뀔 때 대화 기록 완전 초기화\n(현재는 최근 2개만 유지)',
         '6/9~10'),
        ('카메라가 너무 자주 열렸던 문제',
         'AI가 완료 확인할 때도\n카메라를 켜라고 잘못 판단했음',
         'AI 지시 문구 수정 완료\n이번 주 적용됨',
         '완료'),
        ('단계 완료를 AI가 판단하던 문제',
         'AI가 직접 결정하니까\n오판이 잦고 보호자 통제 불가',
         '앱이 키워드로 직접 판단\n이번 주 구조 변경 완료',
         '완료'),
    ],
    widths=[3.5, 4.5, 5, 2]
)
doc.add_paragraph()

# ════════════════════ 5. 남은 일정 ════════════════════
sec(doc, '5', '남은 일정  (6/8 → 6/15 마감)')

img(doc, r'C:\Users\User\Desktop\diagram_gantt.png', width_cm=16.5)

tbl(doc,
    ['날짜', '작업', '목표'],
    [
        ('6/9 (월) ~ 6/10 (화)', 'AI 응답 튜닝', '안내 문구 길이 제한, 단계 혼동 제거'),
        ('6/9 (월) ~ 6/12 (목)', '보호자 앱 UI 개편', '일정 등록·단계 검토 화면 사용성 개선'),
        ('6/11 (수) ~ 6/13 (금)', '실기기 전체 테스트', '실물 기기에서 전체 흐름 시나리오 확인'),
        ('6/13 (금) ~ 6/14 (토)', '버그 수정 · 최종 점검', '테스트에서 나온 문제 수정'),
        ('6/12 (목) ~ 6/15 (일)', '발표 자료 준비', 'PPT 또는 시연 시나리오 작성'),
        ('6/15 (일)', '마감', '제출 완료'),
    ],
    widths=[4, 4.5, 7.5]
)

doc.add_paragraph()
hr(doc)
center(doc, '눈길 팀  |  한이음 ICT 멘토링 클럽  |  2026. 06', size=9, color=C_GRAY)

out = r'C:\Users\User\Desktop\눈길_주간보고서_최종.docx'
doc.save(out)
print(f'저장 완료: {out}')
