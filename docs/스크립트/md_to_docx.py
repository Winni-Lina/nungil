# -*- coding: utf-8 -*-
"""technical_report.md → Word(.docx) 변환 (이미지·표·캡션 포함)"""
import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r'C:\Users\User\Desktop\Nungil Main'
BASE = ROOT  # 이미지 경로(docs/...)는 repo 루트 기준
MD = os.path.join(ROOT, '기술보고서.md')
OUT = os.path.join(ROOT, 'docs', '눈길_Technical_Report.docx')
FONT = '맑은 고딕'

doc = Document()
# 기본 스타일 폰트
style = doc.styles['Normal']
style.font.name = FONT
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

def set_kfont(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    if size: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if color: run.font.color.rgb = color

INLINE = re.compile(r'\*\*(.+?)\*\*|`(.+?)`')
def add_runs(p, text):
    """굵게(**), 코드(`) 인라인 처리"""
    text = text.replace('✅','O').replace('⚠️','△')
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            set_kfont(p.add_run(text[pos:m.start()]))
        if m.group(1) is not None:
            set_kfont(p.add_run(m.group(1)), bold=True)
        else:
            r = p.add_run(m.group(2)); set_kfont(r); r.font.name='Consolas'
            r.font.color.rgb = RGBColor(0xC0,0x30,0x30)
        pos = m.end()
    if pos < len(text):
        set_kfont(p.add_run(text[pos:]))

def img_path(src):
    src = src.replace('%20',' ').replace('/', os.sep)
    if src.startswith('..'+os.sep):
        return os.path.normpath(os.path.join(BASE, src))
    return os.path.normpath(os.path.join(BASE, src))

def add_table(rows):
    cols = len(rows[0])
    t = doc.add_table(rows=len(rows), cols=cols)
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = t.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            txt = row[ci] if ci < len(row) else ''
            txt = txt.replace('<br>','\n')
            add_runs(p, txt)
            for run in p.runs:
                set_kfont(run, size=9, bold=(ri==0))
    doc.add_paragraph()

lines = open(MD, encoding='utf-8').read().split('\n')
i = 0
first_h1 = True
while i < len(lines):
    ln = lines[i].rstrip()
    # 코드블록
    if ln.strip().startswith('```'):
        i += 1; code=[]
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code.append(lines[i]); i += 1
        i += 1
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'),'F2F2F2'); pPr.append(shd)
        r = p.add_run('\n'.join(code)); r.font.name='Consolas'; r.font.size=Pt(9)
        continue
    # 이미지
    m = re.match(r'!\[(.*?)\]\((.+?)\)', ln.strip())
    if m:
        path = img_path(m.group(2))
        if os.path.exists(path):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try: p.add_run().add_picture(path, width=Inches(5.6))
            except Exception as e: add_runs(doc.add_paragraph(), f'[이미지: {path}]')
        i += 1; continue
    # 표 (연속 | 행 수집, 구분선 스킵)
    if ln.strip().startswith('|'):
        block=[]
        while i < len(lines) and lines[i].strip().startswith('|'):
            block.append(lines[i].strip()); i += 1
        rows=[]
        for b in block:
            if re.match(r'^\|[\s:\-\|]+\|?$', b): continue
            cells=[c.strip() for c in b.strip('|').split('|')]
            rows.append(cells)
        if rows: add_table(rows)
        continue
    # 헤딩
    if ln.startswith('#'):
        lvl = len(ln) - len(ln.lstrip('#'))
        txt = ln.lstrip('#').strip()
        if lvl == 1:
            p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            set_kfont(p.add_run(txt), size=20, bold=True, color=RGBColor(0x1A,0x5F,0x9E))
            if first_h1: first_h1=False
        else:
            p = doc.add_paragraph()
            sizes={2:15,3:12.5,4:11}
            set_kfont(p.add_run(txt), size=sizes.get(lvl,11), bold=True,
                      color=RGBColor(0x1A,0x5F,0x9E) if lvl==2 else RGBColor(0x33,0x33,0x33))
            p.paragraph_format.space_before = Pt(10 if lvl==2 else 6)
        i += 1; continue
    # 구분선
    if ln.strip() == '---':
        i += 1; continue
    # 인용
    if ln.strip().startswith('>'):
        txt = ln.strip().lstrip('>').strip()
        if txt:
            p = doc.add_paragraph()
            pPr=p._p.get_or_add_pPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'EEF4FB'); pPr.append(shd)
            p.paragraph_format.left_indent=Inches(0.15)
            add_runs(p, txt)
            for r in p.runs: r.font.size=Pt(9.5)
        i += 1; continue
    # 리스트
    if re.match(r'^\s*[-*]\s+', ln):
        txt = re.sub(r'^\s*[-*]\s+','',ln)
        p = doc.add_paragraph(style='List Bullet')
        add_runs(p, txt)
        i += 1; continue
    if re.match(r'^\s*\d+\.\s+', ln):
        txt = re.sub(r'^\s*\d+\.\s+','',ln)
        p = doc.add_paragraph(style='List Number')
        add_runs(p, txt)
        i += 1; continue
    # <div>/</div> 무시
    if ln.strip() in ('<div align="center">','</div>',''):
        i += 1; continue
    # 일반 문단
    add_runs(doc.add_paragraph(), ln.strip())
    i += 1

doc.save(OUT)
print('saved:', OUT)
